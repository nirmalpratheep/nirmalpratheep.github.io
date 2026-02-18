"""Generate a 1-page professional resume in Word format — AI/ML Research Engineer."""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn


def set_cell_border(cell, **kwargs):
    """Set cell border. Usage: set_cell_border(cell, bottom={"sz": 4, "color": "000000"})"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.find(qn("w:tcBorders"))
    if tcBorders is None:
        from docx.oxml import OxmlElement
        tcBorders = OxmlElement("w:tcBorders")
        tcPr.append(tcBorders)
    for edge, attrs in kwargs.items():
        from docx.oxml import OxmlElement
        element = OxmlElement(f"w:{edge}")
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), str(attrs.get("sz", 4)))
        element.set(qn("w:color"), attrs.get("color", "000000"))
        element.set(qn("w:space"), "0")
        tcBorders.append(element)


def add_styled_paragraph(doc, text, font_size=9, bold=False, color=None,
                         space_after=2, space_before=0, alignment=None, font_name="Calibri"):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.line_spacing = Pt(font_size + 2.5)
    if alignment:
        p.alignment = alignment
    run = p.add_run(text)
    run.font.size = Pt(font_size)
    run.font.name = font_name
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)
    return p


def add_mixed_paragraph(doc, segments, space_after=2, space_before=0, line_spacing=None):
    """Add paragraph with mixed formatting. segments = [(text, font_size, bold, color), ...]"""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    if line_spacing:
        p.paragraph_format.line_spacing = Pt(line_spacing)
    for text, font_size, bold, color in segments:
        run = p.add_run(text)
        run.font.size = Pt(font_size)
        run.font.name = "Calibri"
        run.bold = bold
        if color:
            run.font.color.rgb = RGBColor(*color)
    return p


def add_bullet(doc, text_segments, space_after=1, indent=Inches(0.25)):
    """Add a bullet point with mixed bold/normal text."""
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.line_spacing = Pt(11)
    p.paragraph_format.left_indent = indent
    # Clear default run
    for run in p.runs:
        run.text = ""
    for text, bold in text_segments:
        run = p.add_run(text)
        run.font.size = Pt(8.5)
        run.font.name = "Calibri"
        run.bold = bold
    return p


def add_section_header(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.line_spacing = Pt(12)
    run = p.add_run(text.upper())
    run.font.size = Pt(9.5)
    run.font.name = "Calibri"
    run.bold = True
    run.font.color.rgb = RGBColor(30, 58, 95)
    # Add bottom border
    from docx.oxml import OxmlElement
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "2563EB")
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def generate_resume():
    doc = Document()

    # Page margins - tight for 1-page
    for section in doc.sections:
        section.top_margin = Inches(0.4)
        section.bottom_margin = Inches(0.3)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)

    # --- HEADER ---
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(1)
    run = p.add_run("NIRMAL PRATHEEP NATARAJAN")
    run.font.size = Pt(16)
    run.font.name = "Calibri"
    run.bold = True
    run.font.color.rgb = RGBColor(30, 58, 95)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.space_before = Pt(0)
    run = p.add_run("AI/ML Research Engineer")
    run.font.size = Pt(11)
    run.font.name = "Calibri"
    run.font.color.rgb = RGBColor(37, 99, 235)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.space_before = Pt(1)
    run = p.add_run("nirmalpratheep@gmail.com  |  linkedin.com/in/nirmalpratheep  |  github.com/nirmalpratheep")
    run.font.size = Pt(8.5)
    run.font.name = "Calibri"
    run.font.color.rgb = RGBColor(74, 85, 104)

    # --- SUMMARY ---
    add_section_header(doc, "Summary")
    add_styled_paragraph(
        doc,
        "AI/ML Research Engineer with 13+ years at AMD & Xilinx, specializing in implementing and optimizing "
        "ML research from paper to production. Hands-on experience with LLM pre-training & alignment "
        "(SFT, GRPO, RLHF), deep reinforcement learning for combinatorial optimization, and GPU kernel-level "
        "performance optimization (Triton, Flash Attention, Nsight profiling). Proven track record of translating "
        "research ideas into working systems, designing experiments, and publishing at peer-reviewed conferences.",
        font_size=8.5, space_after=2
    )

    # --- RESEARCH & PUBLICATIONS ---
    add_section_header(doc, "Research & Publications")
    research_bullets = [
        [("Deep RL for FloorPlan Optimization ", True), ("\u2014 GTAC'25 & SPS Tech Conference (Finalist, arXiv pending). ", False),
         ("Formulated FPGA floorplan optimization as RL; GIN on 15M-node netlists; 2% placement QoR gain", False)],
        [("ML-based Delay Prediction ", True), ("\u2014 GTAC'22 AMD Tech Conference (Finalist). ", False),
         ("ML delay models + GNN design complexity analysis with drift detection", False)],
        [("LLM Alignment & Reasoning via RL ", True), ("\u2014 End-to-end pipeline (Baseline \u2192 SFT \u2192 GRPO RL) on Qwen 2.5 Math 1.5B; ", False),
         ("14.2\u00d7 zero-shot accuracy gain", True), (" with systematic ablation", False)],
        [("Adaptive OFDM Pilots ", True), ("\u2014 IEEE WAMICON 2009", False)],
    ]
    for segments in research_bullets:
        add_bullet(doc, segments)

    # --- CORE EXPERTISE ---
    add_section_header(doc, "Core Expertise")
    expertise_items = [
        ("ML Frameworks & Training: ", "PyTorch, HuggingFace, TRL, vLLM, DeepSpeed/FSDP, Flash Attention, mixed-precision"),
        ("GPU & Performance: ", "Triton kernel development, Nsight Systems/Compute profiling, CUDA, kernel-level optimization"),
        ("RL & Agents: ", "Stable Baselines 3, Ray, Gym, LangGraph/LangChain, GNN feature extraction, policy optimization"),
        ("Infrastructure: ", "Kubernetes, Docker, Ray, LSF, W&B, Optuna | Python, C++, Golang"),
    ]
    for label, detail in expertise_items:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(1)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.line_spacing = Pt(11)
        run = p.add_run(label)
        run.font.size = Pt(8.5)
        run.font.name = "Calibri"
        run.bold = True
        run = p.add_run(detail)
        run.font.size = Pt(8.5)
        run.font.name = "Calibri"

    # --- EXPERIENCE ---
    add_section_header(doc, "Professional Experience")

    # AMD header
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.line_spacing = Pt(12)
    run = p.add_run("AMD ")
    run.font.size = Pt(9.5)
    run.font.name = "Calibri"
    run.bold = True
    run = p.add_run("(formerly Xilinx)")
    run.font.size = Pt(8.5)
    run.font.name = "Calibri"
    run = p.add_run("  |  Senior Staff / Research Engineer \u2014 AI/ML & Design Automation")
    run.font.size = Pt(8.5)
    run.font.name = "Calibri"
    run.font.color.rgb = RGBColor(37, 99, 235)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(0)
    run = p.add_run("San Jose, CA  |  2012 \u2013 Present")
    run.font.size = Pt(8)
    run.font.name = "Calibri"
    run.font.color.rgb = RGBColor(113, 128, 150)

    # -- Research & ML Systems --
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.line_spacing = Pt(11)
    run = p.add_run("Research & ML Systems")
    run.font.size = Pt(8.5)
    run.font.name = "Calibri"
    run.bold = True
    run.font.color.rgb = RGBColor(30, 58, 95)

    research_bullets = [
        [("Deep RL for directive optimization: ", True), ("designed environment, reward shaping, GIN feature extraction on 15M-node netlists; 2% placement QoR; ", False), ("GTAC'25 Finalist", True)],
        [("Ray-based distributed training ", True), ("infrastructure with Grid, ASHA, PBT hyperparameter search for systematic experiment management", False)],
        [("ML delay prediction + GNN design complexity models ", True), ("with automated fine-tuning, monitoring, drift detection; ", False), ("GTAC'22 Finalist", True)],
        [("Agentic AI framework ", True), ("using LangGraph/LLMs for autonomous triage with self-correction and Dockerized evaluation", False)],
    ]
    for segments in research_bullets:
        add_bullet(doc, segments)

    # -- Performance Engineering & Infrastructure --
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.line_spacing = Pt(11)
    run = p.add_run("Performance Engineering & Infrastructure")
    run.font.size = Pt(8.5)
    run.font.name = "Calibri"
    run.bold = True
    run.font.color.rgb = RGBColor(30, 58, 95)

    infra_bullets = [
        [("Led 10+ engineer team ", True), ("across global sites for simulation tooling on ", False), ("10nm/7nm/2nm FPGA nodes", True)],
        [("Architected client/server system ", True), ("(Boost Asio + Protobuf) with ", False), ("3x throughput", True), ("; divide-and-conquer parallel processing enabling ", False), ("20x scale", True)],
        [("Graph compression pipeline ", True), ("processing 3.5B datapoints, reducing 1B paths to 500K patterns", False)],
    ]
    for segments in infra_bullets:
        add_bullet(doc, segments)

    # --- RESEARCH ENGINEERING PROJECTS ---
    add_section_header(doc, "Research Engineering Projects")
    projects = [
        ("LLM Alignment & Reasoning RL: ", "Baseline \u2192 SFT \u2192 GRPO RL on Qwen 2.5 Math 1.5B. 14.2\u00d7 accuracy gain. TRL GRPOTrainer + vLLM colocate, dual-GPU Optuna + ASHA"),
        ("1B Seed Model Pre-training: ", "33% throughput gain via custom Triton kernels, fused Flash Attention. Nsight Systems/Compute profiling"),
        ("SmolLM v2 Pre-training: ", "135M params on FineWeb-Edu. ~40k tokens/sec (BF16); loss 11.6 \u2192 0.0015"),
        ("Agentic Coding Pipeline: ", "LangGraph orchestration, AST analysis, iterative self-correction. 90%+ success rate"),
        ("SWE-Agent Benchmark: ", "3-agent eval architecture with Docker test isolation on SWE-bench"),
        ("ImageNet Classifier: ", "ResNet-50, 77.4% Top-1. Also: Tamil BPE Tokenizer, RL Car Navigation, MNIST Architecture Search"),
    ]
    for label, detail in projects:
        add_bullet(doc, [(label, True), (detail, False)], space_after=1)

    # --- EDUCATION & CERTS ---
    add_section_header(doc, "Education & Certifications")
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.line_spacing = Pt(11)
    run = p.add_run("M.Eng ")
    run.font.size = Pt(8.5)
    run.font.name = "Calibri"
    run.bold = True
    run = p.add_run("Electrical Engineering \u2014 University of Cincinnati, 2012   |   ")
    run.font.size = Pt(8.5)
    run.font.name = "Calibri"
    run = p.add_run("B.Eng ")
    run.font.size = Pt(8.5)
    run.font.name = "Calibri"
    run.bold = True
    run = p.add_run("Electronics & Communication \u2014 Anna University, 2007")
    run.font.size = Pt(8.5)
    run.font.name = "Calibri"

    add_styled_paragraph(
        doc,
        "Certifications: Triton Kernel Dev (AMD Instinct GPUs) \u2022 LLM Serving with vLLM & MI300X \u2022 "
        "Agentic Framework (HuggingFace) \u2022 Generative AI with LLMs (DeepLearning.AI) \u2022 "
        "ML Ops (DeepLearning.AI) \u2022 Machine Learning (Stanford) \u2022 Analytics Edge (MITx) \u2022 "
        "Parallel & Distributed Computing (Rice) \u2022 Kubernetes (Udacity) \u2022 Big Data with Spark (Berkeley)",
        font_size=8, space_after=0, space_before=2, color=(113, 128, 150)
    )

    # Honors
    add_styled_paragraph(
        doc,
        "Top 15 \u2014 Innovate India Design Contest (ALTERA, 2007)  |  AMD Elite Mentorship Program",
        font_size=8, space_after=0, space_before=2, color=(113, 128, 150)
    )

    # Save
    output_path = r"c:\Users\ssuga\nirmalp\nirmalpratheep.github.io\Nirmal_Pratheep_Natarajan_Resume.docx"
    doc.save(output_path)
    print(f"Resume saved to: {output_path}")


if __name__ == "__main__":
    generate_resume()
